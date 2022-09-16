#数据报表的生成（Excel+Word）
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

imgname='chart.jpg'
stu=pd.read_excel(r'C:\Users\ASUS\Desktop\MLCode\student.xlsx')
stu.sort_values(by='Score',inplace=True,ascending=False)

plt.bar(stu.Name,stu.Score,color='orange')

plt.title('Score Chart')
plt.xlabel('Name:')
plt.ylabel('Score:')

plt.tight_layout()
plt.savefig(imgname)

docu=Document()
#标题
p=docu.add_heading('班里学生成绩情况：',level=0)
f_stu=stu.iloc[0,:]['Name']#绝对位置去获取
f_score=stu.iloc[0,:]['Score']#分数
p=docu.add_paragraph('班里的第一名是：')
p.add_run(f'{str(f_stu)}分数为：{str(f_score)}').bold=True
#班里情况
p1=docu.add_paragraph(f'有{len(stu.Name)}名学生考试了，总体情况如下：')
table=docu.add_table(rows=len(stu.Name)+1,cols=2)

table.cell(0,0).text='学生姓名:'
table.cell(0,1).text='学生分数:'
for i,(index,row) in enumerate(stu.iterrows()):
    table.cell(i+1,0).text=str(row['Name'])
    table.cell(i+1,1).text=str(row['Score'])
docu.add_picture(imgname)
docu.save('studentsReport.docx')
print('finish')
